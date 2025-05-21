from fastapi import FastAPI, Request
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_UNDERLINE
from io import BytesIO
import json
import os

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],  # Update for production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.post("/generate-doc")
async def generate_doc(request: Request):
    data = await request.json()

    # Safe logging for Unicode characters
    print(json.dumps(data, ensure_ascii=False))

    # Load the template from the template folder
    template_path = os.path.join("templates", "FU_TEMPLATE_Klickovich.docx")
    doc = Document(template_path)

    # Helper to add section headings
    def add_heading(text):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = True
        run.underline = WD_UNDERLINE.SINGLE
        run.font.size = Pt(12)

    # Add 'Other Plans' section
    if "otherPlans" in data and "lines" in data["otherPlans"]:
        add_heading(data["otherPlans"].get("heading", "Other Plans:"))
        for i, line in enumerate(data["otherPlans"]["lines"], start=1):
            doc.add_paragraph(f"{i}. {line}")


    # Add 'Facet, RFA, & ESI/Caudal Injection Notes'
    if "formattedLines" in data:
        add_heading("Facet, RFA, & ESI/Caudal Injection Notes:")
        for line in data["formattedLines"].split("\n"):
            doc.add_paragraph(line)

    # Add follow-up appointment
    if "followUpAppointment" in data:
        add_heading("Follow-up Appointment:")
        doc.add_paragraph(data["followUpAppointment"])

    # Add signature block
    if "signatureLine" in data:
        add_heading("Signature Note:")
        doc.add_paragraph(data["signatureLine"])

    # Add transcription date
    if "dateTranscribed" in data:
        add_heading("Date Transcribed:")
        doc.add_paragraph(data["dateTranscribed"])

    # Generate and return DOCX
    output = BytesIO()
    doc.save(output)
    output.seek(0)

    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": 'attachment; filename="FollowUpNotes.docx"'
        },
    )
